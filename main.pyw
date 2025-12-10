import customtkinter as ctk
import requests
import json
from datetime import date
from tkinter import filedialog
import re
import os
import tkinter as tk
import random

ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("green")

class StoryGenerator:
    def __init__(self):
        self.root = ctk.CTk()
        self.root.title("Story Generator")
        self.root.geometry("800x700")
        self.root.minsize(800, 700)
        self.root.maxsize(800, 700)

        self.load_genres()

        # Main container
        main = ctk.CTkFrame(self.root)
        main.pack(fill="both", expand=True, padx=30, pady=30)

        # Title
        ctk.CTkLabel(main, text="Story Generator", font=ctk.CTkFont(size=26, weight="bold")).pack(pady=(0, 25))

        # Genre
        genre_frame = ctk.CTkFrame(main, fg_color="transparent")
        genre_frame.pack(pady=(5, 10))
        ctk.CTkLabel(genre_frame, text="Select Genre", font=ctk.CTkFont(size=12)).pack(side="left", padx=(0,10))
        self.genre_var = ctk.StringVar(value=self.genres[0] if self.genres else "")
        self.genre_combo = ctk.CTkComboBox(genre_frame, values=self.genres, variable=self.genre_var, width=250, height=30,
                        font=ctk.CTkFont(size=12))
        self.genre_combo.pack(side="left")
        ctk.CTkButton(genre_frame, text="⚙️", command=self.manage_genres, width=30, height=30,
                      font=ctk.CTkFont(size=12)).pack(side="left", padx=(5,0))

        # Word count
        ctk.CTkLabel(main, text="Word Count", font=ctk.CTkFont(size=12)).pack(anchor="w", padx=100)
        self.slider_var = ctk.DoubleVar(value=500)
        slider = ctk.CTkSlider(main, from_=100, to=1500, number_of_steps=140, variable=self.slider_var, width=400)
        slider.pack(pady=(10, 10))
        self.word_label = ctk.CTkLabel(main, text="500 words", font=ctk.CTkFont(size=14, weight="bold"))
        self.word_label.pack(pady=(0, 10))
        slider.configure(command=lambda v: self.word_label.configure(text=f"{int(float(v))} words"))

        # Buttons
        btn_frame = ctk.CTkFrame(main, fg_color="transparent")
        btn_frame.pack(pady=5)
        ctk.CTkButton(btn_frame, text="Generate Story", command=self.generate_story,
                      fg_color="#2364EF", hover_color="#0407D9", width=150, height=30,
                      font=ctk.CTkFont(size=11, weight="bold")).pack(side="left", padx=30)
        self.save_btn = ctk.CTkButton(btn_frame, text="Save Story", command=self.save_story,
                                      fg_color="#2D936C", hover_color="#1E6134", width=150, height=30,
                                      font=ctk.CTkFont(size=11, weight="bold"), state="disabled")
        self.save_btn.pack(side="left", padx=30)

        # Story display
        self.story_text = ctk.CTkTextbox(main, wrap="word", font=ctk.CTkFont(size=16), padx=20, pady=20,
                                         scrollbar_button_color="#555555", scrollbar_button_hover_color="#777777")
        self.story_text.pack(fill="both", expand=True, padx=50, pady=(5, 10))

    def generate_story(self):
        genre = self.genre_var.get()
        words = int(self.slider_var.get())

        prompt = (f"Write a complete, original {genre.lower()} story that is up to {words} words long (count carefully). "
                  "CRITICAL: The story MUST have a proper ending - do not stop mid-sentence or mid-thought. "
                  "Format with multiple paragraphs separated by blank lines (double newlines). Always include a title as the first line, with a blank line after it. "
                  "Start new paragraphs at natural breaks: scene changes, dialogue, shifts in time/location. "
                  "Use **bold** for strong emphasis and *italics* for thoughts or sounds. "
                  "Strong plot, vivid characters, satisfying ending.")

        self.story_text.delete("1.0", "end")
        self.story_text.insert("1.0", "Generating your story...")
        self.root.update()

        try:
            with open("api_key.json", "r") as f:
                data = json.load(f)
                api_key = data["api_key"]
            
            # Use Groq's free API (llama-3.3-70b-versatile model)
            response = requests.post(
                "https://api.groq.com/openai/v1/chat/completions",
                json={
                    "model": "llama-3.3-70b-versatile",
                    "messages": [
                        {"role": "system", "content": "You are a master storyteller with extensive experience in creative writing. Always create completely original stories with unique characters, plots, and settings. Focus on vivid descriptions, emotional depth, and engaging narratives that captivate readers."},
                        {"role": "user", "content": prompt}
                    ],
                    "max_tokens": words * 5,
                    "temperature": 1.7, # Creative; anything less tends to reproduce similar stories
                    "seed": random.randint(0, 1000000)
                },
                headers={"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"},
                timeout=60
            )
            response.raise_for_status()
            raw = response.json()["choices"][0]["message"]["content"].strip()
            story = raw  # Use the complete story as generated by the LLM

            self.story_text.delete("1.0", "end")
            self.display_formatted(story)

            self.current_story = story
            self.current_genre = genre
            self.current_words = words
            self.save_btn.configure(state="normal")

        except Exception as e:
            self.story_text.delete("1.0", "end")
            self.story_text.insert("1.0", f"Failed to generate:\n\n{str(e)}")

    def display_formatted(self, text):
        # Split by double newlines (paragraph breaks)
        paragraphs = text.split('\n\n')
        
        for i, paragraph in enumerate(paragraphs):
            paragraph = paragraph.strip()
            if not paragraph:
                continue
                
            # Split paragraph by markdown formatting
            parts = re.split(r'(\*\*.*?\*\*|__.*?__|\*.*?\*|_\*?._\*?)', paragraph)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    self.story_text.insert("end", part[2:-2], "bold")
                elif part.startswith("__") and part.endswith("__"):
                    self.story_text.insert("end", part[2:-2], "bold")
                elif (part.startswith("*") and part.endswith("*") and not part.startswith("**")) or (part.startswith("_") and part.endswith("_") and not part.startswith("__")):
                    self.story_text.insert("end", part[1:-1], "italic")
                else:
                    self.story_text.insert("end", part)
            
            # Add double newline after paragraph for spacing
            if i < len(paragraphs) - 1:
                self.story_text.insert("end", "\n\n")

    def save_story(self):
        today = date.today().strftime("%Y-%m-%d")

        # Extract title from first line of story
        lines = self.current_story.split('\n', 1)
        if len(lines) > 1 and lines[0].strip():
            title = lines[0].strip()
            content = lines[1].strip()
        else:
            title = f"{self.current_genre} Story"
            content = self.current_story

        # Sanitize title for filename
        import re
        safe_title = re.sub(r'[^\w\s-]', '', title).replace(' ', '_')
        filename = f"{safe_title}_{self.current_words}w_{today}.md"

        path = filedialog.asksaveasfilename(
            initialfile=filename,
            defaultextension=".md",
            filetypes=[("DOCX", "*.docx"), ("Markdown", "*.md"), ("RTF", "*.rtf"), ("PDF", "*.pdf"), ("Text", "*.txt")],
            title="Save Story"
        )

        ext = os.path.splitext(path)[1].lower()

        if ext == '.md':
            content_full = f"# {title}\n\n**{self.current_words} words**\n\n---\n\n{content}"
        elif ext == '.txt':
            content_full = self.current_story
        elif ext == '.rtf':
            content_full = self.generate_rtf(content, title)
        elif ext == '.docx':
            self.save_as_docx(path, content, title)
            self.save_btn.configure(text="Saved!")
            self.root.after(2000, lambda: self.save_btn.configure(text="Save Story"))
            return
        elif ext == '.pdf':
            self.save_as_pdf(path, content, title)
            self.save_btn.configure(text="Saved!")
            self.root.after(2000, lambda: self.save_btn.configure(text="Save Story"))
            return
        else:
            content_full = self.current_story

        try:
            with open(path, "w", encoding="utf-8") as f:
                f.write(content)
            self.save_btn.configure(text="Saved!")
            self.root.after(2000, lambda: self.save_btn.configure(text="Save Story"))
        except Exception as e:
            self.story_text.insert("end", f"\n\nSave failed: {e}")

    def markdown_to_html(self, text):
        import re
        # Convert bold **text** or __text__ to <b>text</b>
        text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
        text = re.sub(r'__(.*?)__', r'<b>\1</b>', text)
        # Convert italics *text* or _text_ to <i>text</i>
        text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text)
        text = re.sub(r'_(.*?)_', r'<i>\1</i>', text)
        return text

    def generate_rtf(self, text, title):
        # RTF generation with bold and italic formatting
        rtf = r"{\rtf1\ansi\deff0{\fonttbl{\f0\fnil\fcharset0 Arial;}}\f0\fs24 "
        rtf += r'\b ' + title + r'\b0\par '
        rtf += r'\i ' + f"{self.current_words} words" + r'\i0\par \par '
        paragraphs = text.split('\n\n')
        for para in paragraphs:
            if para.strip():
                parts = re.split(r'(\*\*.*?\*\*|__.*?__|\*.*?\*|_\*?._\*?)', para)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        rtf += r'\b ' + part[2:-2] + r'\b0 '
                    elif part.startswith("__") and part.endswith("__"):
                        rtf += r'\b ' + part[2:-2] + r'\b0 '
                    elif (part.startswith("*") and part.endswith("*") and not part.startswith("**")) or (part.startswith("_") and part.endswith("_") and not part.startswith("__")):
                        rtf += r'\i ' + part[1:-1] + r'\i0 '
                    else:
                        rtf += part + ' '
                rtf += r'\par '
        rtf += r"}"
        return rtf

    def save_as_docx(self, path, text, title):
        try:
            from docx import Document
            doc = Document()
            doc.add_heading(title, 0)
            doc.add_paragraph(f"{self.current_words} words").italic = True
            paragraphs = text.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    p = doc.add_paragraph()
                    parts = re.split(r'(\*\*.*?\*\*|__.*?__|\*.*?\*|_\*?._\*?)', paragraph)
                    for part in parts:
                        if part.startswith("**") and part.endswith("**"):
                            p.add_run(part[2:-2]).bold = True
                        elif part.startswith("__") and part.endswith("__"):
                            p.add_run(part[2:-2]).bold = True
                        elif (part.startswith("*") and part.endswith("*") and not part.startswith("**")) or (part.startswith("_") and part.endswith("_") and not part.startswith("__")):
                            p.add_run(part[1:-1]).italic = True
                        else:
                            p.add_run(part)
            doc.save(path)
        except ImportError:
            raise Exception("python-docx not installed. Install with: pip install python-docx")

    def save_as_pdf(self, path, text, title):
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph

            doc = SimpleDocTemplate(path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            story.append(Paragraph(f"<b>{title}</b>", styles['Heading1']))
            story.append(Paragraph(f"<i>{self.current_words} words</i>", styles['Normal']))
            text_html = self.markdown_to_html(text)
            # Replace double newlines with paragraph breaks, single with <br/>
            paragraphs = text_html.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    story.append(Paragraph(para.replace('\n', '<br/>'), styles['Normal']))
            doc.build(story)
        except ImportError:
            raise Exception("reportlab not installed. Install with: pip install reportlab")

    def run(self):
        self.root.mainloop()

    def load_genres(self):
        try:
            with open("genres.json", "r") as f:
                self.genres = json.load(f)
        except FileNotFoundError:
            self.genres = ["Horror", "Thriller", "Mystery", "Fantasy", "Science Fiction",
                           "Romance", "Western", "Adventure", "Comedy", "Drama",
                           "Action","Time Travel","Post-Apocalyptic","Magical Realism"]
            self.save_genres()

    def save_genres(self):
        with open("genres.json", "w") as f:
            json.dump(self.genres, f, indent=2)

    def manage_genres(self):
        window = ctk.CTkToplevel(self.root)
        window.title("Manage Genres")
        window.geometry("400x300")
        window.attributes("-topmost", True)  # Always on top
        
        # Center the window on the main window
        window.update_idletasks()
        main_x = self.root.winfo_x()
        main_y = self.root.winfo_y()
        main_width = self.root.winfo_width()
        main_height = self.root.winfo_height()
        popup_width = 400
        popup_height = 300
        x = main_x + (main_width - popup_width) // 2
        y = main_y + (main_height - popup_height) // 2
        window.geometry(f"400x300+{x}+{y}")
        
        # Listbox with scrollbar
        list_frame = tk.Frame(window, bg="#2b2b2b")
        list_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        scrollbar = tk.Scrollbar(list_frame, orient="vertical", bg="#555555", activebackground="#777777", troughcolor="#2b2b2b")
        self.genre_listbox = tk.Listbox(list_frame, selectmode=tk.SINGLE, bg="#2b2b2b", fg="white", selectbackground="#1f6aa5", font=("Arial", 12), yscrollcommand=scrollbar.set, highlightthickness=0)
        scrollbar.config(command=self.genre_listbox.yview)
        
        self.genre_listbox.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        for genre in self.genres:
            self.genre_listbox.insert(tk.END, genre)
        
        # Control frame
        control_frame = ctk.CTkFrame(window, fg_color="transparent")
        control_frame.pack(fill="x", padx=10, pady=(0,10))
        
        self.new_genre_entry = ctk.CTkEntry(control_frame, placeholder_text="New genre")
        self.new_genre_entry.pack(side="left", fill="x", expand=True)
        
        ctk.CTkButton(control_frame, text="Add", command=self.add_genre).pack(side="left", padx=(5,0))
        ctk.CTkButton(control_frame, text="Delete", command=self.delete_genre, fg_color="#EF233C").pack(side="left", padx=(5,0))

    def add_genre(self):
        new_genre = self.new_genre_entry.get().strip()
        if new_genre and new_genre not in self.genres:
            self.genres.append(new_genre)
            self.save_genres()
            self.genre_listbox.insert(tk.END, new_genre)
            self.genre_combo.configure(values=self.genres)
            self.new_genre_entry.delete(0, tk.END)

    def delete_genre(self):
        selected = self.genre_listbox.curselection()
        if selected:
            index = selected[0]
            del self.genres[index]
            self.save_genres()
            self.genre_listbox.delete(index)
            self.genre_combo.configure(values=self.genres)
            if self.genres:
                self.genre_var.set(self.genres[0])
            else:
                self.genre_var.set("")

if __name__ == "__main__":
    app = StoryGenerator()
    app.run()